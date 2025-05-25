import os.path
import re
from os import remove
from typing import List
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Cm, Pt

from database import Documents, AnotherDoc, BuilderDoc, Reglaments, select_data

docs: List[str] = ["ГОСТ", "ТР ТС", "ТР ЕАЭС", "СП", "СанПиН", "СниП", "СП", "РД", "ОСТ", "ГН", "ИСО", "НПБ"]
terms = "|".join(docs)
pattern = fr'(?:{terms})\s(?:[A-ZА-Я]+\s)?\d+(?:[/.-]\d+)?(?:[/.-]\d+)?(?:[/.-]\d+)?'


def remove_duplicates(array: List[str]) -> List[str]:
    """Наивный алгоритм удаления дубликатов"""
    length = len(array)
    i = 0
    while i < length:
        found = False
        for k in range(i+1, length):
            if array[k] == array[i]:
                found = True
                break
        if not found:
            i += 1
            continue
        else:
            array.remove(array[i])
        length -= 1
    return array


def writer_table(data: List[tuple]) -> None:
    """Функция для создания файла и записи таблицы"""
    filename: str = 'Таблица с ГОСТ.docx'
    bad_gosts = ["Утратил силу в РФ", "Заменен", "Отменен", "Срок действия истек", "Отловили ошибку", "Неизвестен"]
    data_table = [("Обозначение документа", "Наименование документа")]
    data_table.extend(data)
    p = os.path.abspath(filename)
    if os.path.isfile(p):
        remove(p)
    doc = Document()
    table = doc.add_table(rows=len(data_table), cols=2)
    table.style = "Table Grid"
    for row_idx, row_data in enumerate(data_table):
        row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if row_idx == 0:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(14)
                        run.bold = True
                    elif [s for s in bad_gosts if s in cell_data]: # делать вхождение подстроки
                        run.bold = True
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    else:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(14)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = Pt(1)
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                if row_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    column_widths = [Cm(5), Cm(17)]
    for col_idx, width in enumerate(column_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width
    doc.save(filename)


def prepare_data(ready: List[str]) -> List[tuple]:
    """Функция для подготовки данных перед записью в файл"""
    data_to_file = []
    for gost in ready:
        if gost[:5] in ["ТР ТС", "ТР ЕА"]:
            data: tuple = select_data(Reglaments, gost)
            data_to_file.append(data)
        elif gost[:5] in ["СанПи", "СНиП "] or gost[:2] == "СП":
            data: tuple = select_data(BuilderDoc, gost)
            data_to_file.append(data)
        elif gost[:4] == "ГОСТ":
            data: tuple = select_data(Documents, gost)
            data_to_file.append(data)
        else:
            data: tuple = select_data(AnotherDoc, gost)
            data_to_file.append(data)
    return data_to_file


def extract_gosts(file_path: str) -> None:
    """Функция для чтения файла"""
    doc = Document(file_path)
    ready_docs = []
    for row in doc.paragraphs:
        gosts = re.findall(pattern, row.text)
        ready_docs += gosts

    ready: List[str] = remove_duplicates(ready_docs)
    data: List[tuple] = prepare_data(ready)
    sorted_data: List[tuple] = sorted(data)
    writer_table(sorted_data)
